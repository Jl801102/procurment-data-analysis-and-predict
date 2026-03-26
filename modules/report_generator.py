# modules/report_generator.py
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

class ReportGenerator:
    @staticmethod
    def generate_word_report(strategies, analysis_summary, company_name="XX公司"):
        doc = Document()
        
        # 设置中文字体（避免显示方块）
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5)
        for hs in ['Heading 1', 'Heading 2', 'Heading 3']:
            try:
                h_style = doc.styles[hs]
                h_style.font.name = '宋体'
                h_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except KeyError:
                pass

        # 标题
        title = doc.add_heading('采购数据分析与降本策略报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"报告主体：{company_name}")
        doc.add_paragraph(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
        doc.add_paragraph(f"报告版本：V1.0（系统自动生成）")
        doc.add_page_break()

        # 执行摘要
        doc.add_heading('一、执行摘要', level=1)
        total_saving = analysis_summary.get('总降本金额', 0)
        total_spend = analysis_summary.get('总采购额', 0)
        # 确保是数值
        if not isinstance(total_saving, (int, float)):
            total_saving = 0
        if not isinstance(total_spend, (int, float)):
            total_spend = 0
        saving_percent = (total_saving / total_spend * 100) if total_spend else 0
        doc.add_paragraph(f"本报告基于系统采购数据分析，共识别出 {len(strategies)} 项降本机会，预计年化节省金额约 {total_saving:,.2f} 元，占采购总额的 {saving_percent:.2f}%。")

        # 关键数据一览
        doc.add_heading('二、关键数据一览', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = '指标'
        header_cells[1].text = '数值'
        for key, value in analysis_summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            if isinstance(value, (int, float)):
                if '金额' in key or '降本' in key or '采购额' in key:
                    row_cells[1].text = f"{value:,.2f} 元"
                elif '比例' in key:
                    row_cells[1].text = f"{value:.2f}%"
                else:
                    row_cells[1].text = f"{value:.2f}"
            else:
                row_cells[1].text = str(value)

        # 降本策略详情
        doc.add_heading('三、降本策略详情', level=1)
        for i, strategy in enumerate(strategies, 1):
            doc.add_heading(f"策略 {i}：{strategy.get('名称', '')}", level=2)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            rows_data = [
                ('预计节省金额', f"{strategy.get('金额', 0):,.2f} 元"),
                ('节省比例', f"{strategy.get('比例', 0):.2%}"),
                ('实施难度', strategy.get('难度', '中')),
                ('优先级', strategy.get('优先级', '中'))
            ]
            for j, (label, value) in enumerate(rows_data):
                info_table.rows[j].cells[0].text = label
                info_table.rows[j].cells[1].text = value
            doc.add_paragraph("行动步骤：")
            for step in strategy.get('步骤', []):
                doc.add_paragraph(step, style='List Bullet')
            if strategy.get('风险提示'):
                risk_para = doc.add_paragraph()
                risk_para.add_run("⚠️ 风险提示：").bold = True
                risk_para.add_run(strategy.get('风险提示'))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
